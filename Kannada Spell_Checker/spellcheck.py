# spellcheck.py 
from trie import Trie  
from Dictionary_load import read_categorized_file 
import Levenshtein as lev
import tkinter as tk
from tkinter import scrolledtext, messagebox, ttk
import tkinter.font as tkFont
from ttkthemes import ThemedTk
import time
from reportlab.pdfgen import canvas
from docx import Document
from tkinter import filedialog
from functools import lru_cache


ignored_words = set()  # Global set to store ignored words


def setup_trie():
    # Load sub-dictionaries from the categorized file
    sub_dictionaries = read_categorized_file("categorized.txt")
    # Initialize and populate the Trie
    trie = Trie()
    for category, words in sub_dictionaries.items():
        for word in words:
            trie.insert(word, category)
    return trie

def if_not_kannada(word):
    # Check each character to see if it is within the Kannada unicode range
    for char in word:
        if not ('\u0C80' <= char <= '\u0CFF'):
            return False  # If any character is not Kannada, return False
    return True  # If all characters are Kannada, return True

def is_kannada_number(word):
    kannada_numbers = ['೦', '೧', '೨', '೩', '೪', '೫', '೬', '೭', '೮', '೯']
    return all(char in kannada_numbers for char in word)


def get_all_suffixes(paradigm_tables):
    """ Retrieve all suffixes from the paradigm tables. """
    return [suffix for sublist in paradigm_tables.values() for suffix in sublist]

def find_suffixes(word, all_suffixes):
    """ Identify and return valid suffixes that the word ends with. """
    for suffix in all_suffixes:
        if word.endswith(suffix):
            return suffix
    return None

def generate_suggestions(root, suffixes):
    """ Generate potential valid words by attaching suffixes to the root. """
    return [root + suffix for suffix in suffixes]

def spell_check():
    global trie
    start_time = time.time()
    textbox.config(state=tk.DISABLED)
    input_text = textbox.get("1.0", "end-1c").strip()
    textbox.tag_remove('misspelled', '1.0', tk.END)

    # Clear the output text before running the spell check
    output_text.delete('1.0', tk.END)

    if not input_text:
        messagebox.showinfo("Error", "Please enter text to check.")
        return

    words = input_text.split()
    misspelled_count = 0

    for word in words:
        if word in ignored_words:  # Skip ignored words
            continue

        if not if_not_kannada(word) or is_kannada_number(word):
            continue

        is_misspelled = process_misspelled_word(word, trie, paradigm_tables)
        if is_misspelled:
            mark_word_as_misspelled(word)
            misspelled_count += 1
        
    duration = time.time() - start_time
    status_var.set(f"Spell Check Complete. Time taken: {duration:.2f} seconds.")
    if misspelled_count > 0:
        output_text.insert(tk.END, f"Errors found: {misspelled_count}.")
    else:
        output_text.insert(tk.END, "No Errors found!!")
    textbox.config(state=tk.NORMAL)
    textbox.focus_set()

    
def process_misspelled_word(word, trie, paradigm_tables):
    # Check if the whole word is valid
    exists, _ = trie.search(word)
    if exists:
        return False  # Word is valid

    suffix = find_suffix(word, paradigm_tables)
    root = word[:-len(suffix)] if suffix else word

    # If root and suffix combination is valid, the word is not misspelled
    if suffix and trie.search(root)[0]:
        return False

    # Word is misspelled
    return True


# Function to find a suffix from the word
def find_suffix(word, paradigm_tables):
    all_suffixes = [suffix for sublist in paradigm_tables.values() for suffix in sublist]
    for suffix in all_suffixes:
        if word.endswith(suffix):
            return suffix
    return ''

# Generate similar words based on similarity to the root
def generate_similar_words(root, trie, max_suggestions=8):
    all_words = trie.get_all_words()
    similar_words_with_distances = [
        (word, lev.distance(root, word)) for word in all_words
    ]
    
    similar_words_with_distances.sort(key=lambda x: x[1])
    return [word for word, _ in similar_words_with_distances[:max_suggestions]]


def display_suggestions(event):
    try:
        index = textbox.index(f"@{event.x},{event.y}")
        tagged_ranges = textbox.tag_ranges('misspelled')

        for start, end in zip(tagged_ranges[0::2], tagged_ranges[1::2]):
            if textbox.compare(start, "<=", index) and textbox.compare(index, "<=", end):
                word = textbox.get(start, end)

                suffix = find_suffix(word, paradigm_tables)
                root_word = word[:-len(suffix)] if suffix else word

                # Generate suggestions directly within the function
                similar_roots = generate_similar_words(root_word, trie, max_suggestions=8)
                suggestions = [similar_root + suffix for similar_root in similar_roots] if suffix else similar_roots

                suggestion_menu = tk.Menu(root, tearoff=0)  # Reference the main window correctly
                suggestion_menu.add_command(label="Ignore word", command=lambda s=start, e=end: ignore_word(s, e))

                if suggestions:
                    for suggestion in suggestions:
                        suggestion_menu.add_command(
                            label=suggestion,
                            command=lambda s=suggestion, start=start, end=end: replace_word(start, end, s)
                        )
                else:
                    suggestion_menu.add_command(label="No valid suggestions available", command=lambda: messagebox.showinfo("No Suggestions", "No valid suggestions available."))

                suggestion_menu.tk_popup(event.x_root, event.y_root)
                break
    except Exception as e:
        print(f"Error in display_suggestions: {e}")

#####################Testing#############################
def generate_suggestions_for_word(word, trie, paradigm_tables):
    try:
        root_word, suffix = analyze_word(word, paradigm_tables)
        similar_roots = generate_similar_words(root_word, trie, max_suggestions=8)
        suggestions = [similar_root + suffix for similar_root in similar_roots] if suffix else similar_roots
        return suggestions
    except Exception as e:
        print(f"Error in generate_suggestions_for_word: {e}")
        return []
def analyze_word(word, paradigm_tables):
    suffix = find_suffix(word, paradigm_tables)
    root_word = word[:-len(suffix)] if suffix else word
    return root_word, suffix


def mark_word_as_misspelled(word):
    start_index = textbox.search(word, '1.0', tk.END)
    if start_index:
        end_index = f"{start_index.split('.')[0]}.{int(start_index.split('.')[1]) + len(word)}"
        textbox.tag_add('misspelled', start_index, end_index)

    textbox.update()



def replace_word(start, end, replacement):
    """
    Replaces the text between 'start' and 'end' indices in the textbox with 'replacement'.
    """
    textbox.delete(start, end)  # Delete the current misspelled word
    textbox.insert(start, replacement)  # Insert the selected suggestion

    textbox.tag_remove('misspelled', start, f"{start} + {len(replacement)}c")

def ignore_word(start, end):
    word = textbox.get(start, end)
    ignored_words.add(word)  # Add the word to the set of ignored words
    textbox.tag_remove('misspelled', start, end)

def save_ignored_words(filename):
    if ignored_words:  # Only create the file if there are words to ignore
        with open(filename, 'w') as f:
            for word in ignored_words:
                f.write(f"{word}\n")
        print(f"Ignored words saved to {filename}")  # Debug output
    else:
        print("No words to ignore; no file saved.")


def load_ignored_words(filename):
    try:
        with open(filename, 'r') as f:
            ignored_words.clear()  # Clear existing ignored words before loading new ones
            for line in f:
                ignored_words.add(line.strip())
            print(f"Ignored words loaded from {filename}")
    except FileNotFoundError:
        print(f"No ignored words file found at {filename}")


def save_as_word(text, filename="output.docx"):
    doc = Document()
    for line in text.split("\n"):
        doc.add_paragraph(line)
    doc.save(filename)
def save_file():
    text = textbox.get("1.0", "end-1c")
    file_path = filedialog.asksaveasfilename(
        defaultextension=".pdf",
        filetypes=[("PDF files", "*.pdf"), ("Word documents", "*.docx"), ("Text files", "*.txt")]
    )
    if file_path:
        if file_path.endswith('.docx'):
            save_as_word(text, filename=file_path)
        elif file_path.endswith('.txt'):
            with open(file_path, 'w') as file:
                file.write(text)
        # Save ignored words in a separate file
        ignored_file_path = file_path + ".ignored"
        save_ignored_words(ignored_file_path)
        messagebox.showinfo("Save File", "File and ignored words have been saved successfully.")
def open_file():
    file_path = filedialog.askopenfilename(
        filetypes=[("Text files", "*.txt"), ("Word documents", "*.docx")]
    )
    if file_path:
        # Clear the current content
        textbox.delete("1.0", tk.END)
        
        if file_path.endswith('.txt'):
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
                textbox.insert("1.0", text)
        elif file_path.endswith('.docx'):
            # If the file is a Word document, we need a different method to read it
            from docx import Document
            doc = Document(file_path)
            full_text = []
            for para in doc.paragraphs:
                full_text.append(para.text)
            textbox.insert("1.0", '\n'.join(full_text))

        # Attempt to load ignored words
        ignored_file_path = file_path + ".ignored"
        load_ignored_words(ignored_file_path)
        messagebox.showinfo("Open File", "File and ignored words have been loaded successfully.")
        print(f"Opened and loaded text from {file_path}")  # Debug statement
    else:
        print("No file selected")  # Debug statement
 
def save_file():
    text = textbox.get("1.0", "end-1c")  # Get text from textbox
    # Ask user for the location and name of the file to save
    file_path = filedialog.asksaveasfilename(
        defaultextension=".pdf",
        filetypes=[("PDF files", "*.pdf"), ("Word documents", "*.docx")],
    )
    if not file_path:  # If the user cancels the dialog, file_path will be None or ''
        return

    if file_path.endswith('.docx'):
        save_as_word(text, filename=file_path)
        messagebox.showinfo("Save File", "File has been saved as Word document successfully.")



# GUI Setup
root = ThemedTk()
root.set_theme("winxpblue")

#Width x height:
root.geometry("950x650")
root.title("Kannada Spell Check Tool")

style = ttk.Style(root)

style.configure("Custom.TFrame", background='grey')
style.configure("TButton", font=('Helvetica', 12), borderwidth='4')
style.configure("TLabel", font=('Helvetica', 12), background='light grey')

input_frame = ttk.Frame(root, padding="10")
input_frame.pack(fill='both', expand=True)

output_frame = ttk.Frame(root, padding="10")
output_frame.pack(fill='both', expand=True)
input_frame.configure(style="Custom.TFrame")
output_frame.configure(style="Custom.TFrame")



status_frame = ttk.Frame(root, padding="3")
status_frame.pack(fill='x', side='bottom')
status_frame.configure(style="Custom.TFrame")


textbox = scrolledtext.ScrolledText(input_frame, height=10, wrap=tk.WORD)
textbox.pack(pady=10, fill='both', expand=True)
textbox.bind('<Button-3>', display_suggestions)

customFont = tkFont.Font(family="Lohit Kannada", size=11) 
textbox.configure(font=customFont)



# Create a frame to contain the buttons side by side
button_frame = ttk.Frame(input_frame)
button_frame.pack(pady=5)




save_button = ttk.Button(button_frame, text="Save file", command=save_file)
save_button.pack(side='left', padx=5)

# Now pack the buttons into the button_frame side by side
check_button = ttk.Button(button_frame, text="Check Spelling", command=spell_check)
check_button.pack(side='left', padx=5)

open_button = ttk.Button(button_frame, text="Open File", command=open_file)  # New button for opening files
open_button.pack(side='left', padx=5)

# Spacer label with desired background color using tk.Label
spacer = tk.Label(button_frame, background='grey', borderwidth=1, relief='sunken')

output_text = scrolledtext.ScrolledText(output_frame, height=5, wrap=tk.WORD)
output_text.pack(pady=10, fill='both', expand=True)


status_var = tk.StringVar()
status_label = ttk.Label(status_frame, textvariable=status_var, background='light grey')
status_label.pack(side='left')


trie = setup_trie()


textbox.tag_configure('misspelled', underline=True, foreground='red')

paradigm_tables = {
    "1": ["ತ್ತಿದ್ದಳು", "ತ್ತಿದ್ದನು", "ತ್ತಿದ್ದಾರೆ", "ತ್ತೀಯ", "ತ್ತಾರೆ"],
    "2": ["ಗಳನ್ನು", "ಗಳಲ್ಲಿ", "ಗಳ"],
    "3": ["ದ್ದನು", "ದ್ದಳು", "ದ್ದರು"],
    "4": ["ದ", "ದಲ್ಲಿ", "ದಿಂದ"],
     }

# Keyboard Layout - Adjusted for more characters per row
kannada_characters = [
    ['ಅ', 'ಆ', 'ಇ', 'ಈ', 'ಉ', 'ಊ', 'ಋ', 'ಎ', 'ಏ', 'ಐ', 'ಒ', 'ಓ', 'ಔ', 'ಕ', 'ಖ', 'ಗ', 'ಘ', 'ಙ', 'ಚ', 'ಛ'],
    ['ಜ', 'ಝ', 'ಞ', 'ಟ', 'ಠ', 'ಡ', 'ಢ', 'ಣ', 'ತ', 'ಥ', 'ದ', 'ಧ', 'ನ', 'ಪ', 'ಫ', 'ಬ', 'ಭ', 'ಮ', 'ಯ', 'ರ', 'ಱ'],
    ['ಲ', 'ವ', 'ಶ', 'ಷ', 'ಸ', 'ಹ', 'ಳ','ಾ', 'ೕ', 'ೖ', 'ೊ', 'ೋ', 'ೌ', '್', 'ೀ', 'ು', 'ೂ', 'ೃ', 'ೄ', 'ೆ', 'ೇ', 'ೈ'],
    ["್ಕ", "್ಖ", "್ಗ", "್ಘ", "್ಙ", "್ಚ", "್ಛ", "್ಜ", "್ಝ", "್ಞ", "್ಟ", "್ಠ", "್ಡ", "್ಢ", "್ಣ", "್ತ", "್ಥ", "್ದ", "್ಧ", "್ನ"],
    ["್ಪ", "್ಫ", "್ಬ", "್ಭ", "್ಮ", "್ಯ", "್ರ", "್ಱ", "್ಲ", '್ವ', "್ಶ", "್ಷ", "್ಸ", "್ಹ", "್ಳ", "್಴", "್ಜ್ಞ", '೦', '೧', '೨', '೩', '೪', '೫', '೬', '೭', '೮', '೯'],
]

# Function to insert characters
def insert_character(character):
    textbox.insert(tk.INSERT, character)

# Create the keyboard UI
keyboard_frame = ttk.Frame(root, padding="3")
keyboard_frame.pack(fill='both', side='top', expand=True)
keyboard_frame.configure(style="Custom.TFrame")

status_var = tk.StringVar()
status_label = ttk.Label(status_frame, textvariable=status_var, background='light grey')
status_label.pack(side='left')

# Define the left padding for each row to align the keys
left_paddings = [5,10,15,20,25]  # Decrease the initial padding to move the keyboard to the left

for i, row in enumerate(kannada_characters):
    row_frame = ttk.Frame(keyboard_frame)
    row_frame.configure(style="Custom.TFrame")

    row_frame.pack(fill='x', side='top', padx=(left_paddings[i], 0))
    
    # The left spacer is adjusted for each row to create the staggered effect
    left_spacer = ttk.Frame(row_frame, width=7)  # Reduced width for a more compact keyboard
    left_spacer.pack(side='left', fill='x', expand=True)
    
    for char in row:
        # Adjust button width to fit more characters
        char_button = ttk.Button(row_frame, text=char, width=2.8, command=lambda c=char: insert_character(c))  # Reduced button width
        char_button.pack(side='left', padx=1, pady=1)
    
    # The right spacer is added to ensure the row content is pushed to the left
    right_spacer = ttk.Frame(row_frame, width=7)  # Reduced width for a more compact keyboard
    right_spacer.pack(side='left', fill='x', expand=True)

root.mainloop()